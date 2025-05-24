import os, shutil
from typing import Optional
import fitz
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LC_Document
from langchain_google_genai.embeddings import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS


# Configurable roots
UPLOAD_ROOT = "uploads"
INDEX_PATH  = "faiss_indices"
MAX_PAGES_PER_FILE = 1000

def rebuild_faiss_index(session_id: str) -> Optional[FAISS]:
    """ Rebuild the FAISS index from all files in uploads/{session_id} """
    embedding     = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    splitter      = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    all_docs = []
    user_dir = os.path.join(UPLOAD_ROOT, session_id)
    if not os.path.isdir(user_dir):
        return None

    for fname in os.listdir(user_dir):
        path = os.path.join(user_dir, fname)
        ext  = os.path.splitext(fname)[1].lower()
        text = ""

        if ext == ".pdf":
            doc = fitz.open(path)
            if len(doc) > MAX_PAGES_PER_FILE:
                doc.close()
                continue
            for page in doc:
                text += page.get_text()
            doc.close()

        elif ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        elif ext in {".doc", ".docx"}:
            docx = DocxDocument(path)
            if len(docx.paragraphs) > MAX_PAGES_PER_FILE * 3:
                continue
            for p in docx.paragraphs:
                text += p.text + "\n"

        if not text.strip():
            continue

        for chunk in splitter.split_text(text):
            all_docs.append(LC_Document(page_content=chunk, metadata={"source": fname}))

    if not all_docs:
        return None

    # wipe old index
    if os.path.isdir(INDEX_PATH):
        shutil.rmtree(INDEX_PATH)

    vs = FAISS.from_documents(all_docs, embedding)
    vs.save_local(INDEX_PATH)
    return vs
